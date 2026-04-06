"""
VKRDocument — класс для работы с ВКР документом по ГОСТ.

Использование:
    from vkr_docx import VKRDocument
    vkr = VKRDocument("ВКР 0.2.7.docx")
    vkr.add_heading1("2 ПРОЕКТИРОВАНИЕ ВЕБ-ПРИЛОЖЕНИЯ")
    vkr.add_body_text("Текст параграфа...")
    vkr.add_figure("image.png", "Рисунок 1 — Описание")
    vkr.add_crossrefs()
    vkr.update_toc()
    vkr.save()
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image
import copy
import random
import string
import re
import os

from .config import load_config


def _gen_bookmark_id():
    """Случайный 12-символьный ID в стиле Google Docs."""
    return "".join(random.choices(string.ascii_lowercase + string.digits, k=12))


class VKRDocument:
    """Класс для работы с ВКР документом по ГОСТ."""

    def __init__(self, path: str, config: dict = None):
        self.path = os.path.abspath(path)
        self.doc = Document(self.path)
        self.cfg = config or load_config()
        self._version = self._parse_version()
        self._bm_counter = self._get_max_bookmark_id() + 1

    # --- Конфиг-хелперы ---

    @property
    def _font(self):
        return self.cfg.get("font", {}).get("name", "Times New Roman")

    @property
    def _font_size(self):
        return self.cfg.get("font", {}).get("size", 14)

    @property
    def _font_size_table(self):
        return self.cfg.get("font", {}).get("size_table", 12)

    @property
    def _line_spacing(self):
        return self.cfg.get("paragraph", {}).get("line_spacing_twips", 360)

    @property
    def _first_line_indent(self):
        return self.cfg.get("paragraph", {}).get("first_line_indent_twips", 709)

    @property
    def _max_img_w(self):
        return self.cfg.get("image", {}).get("max_width_cm", 15.0)

    @property
    def _max_img_h(self):
        return self.cfg.get("image", {}).get("max_height_cm", 23.0)

    @property
    def _screenshot_w(self):
        return self.cfg.get("image", {}).get("screenshot_width_cm", 8.0)

    @property
    def _side_by_side_w(self):
        return self.cfg.get("image", {}).get("side_by_side_width_cm", 7.0)

    @property
    def _working_dir(self):
        return self.cfg.get("working_dir", os.path.dirname(self.path))

    # ============================================================
    # Версионирование
    # ============================================================

    def _parse_version(self) -> list:
        """Парсит версию из имени файла ВКР X.Y.Z.

        Схема:
          0.A.B — в работе, A=глава, B=правки (0..999)
          1.0.0 — ВКР полностью готова
          1.0.N — правки после готовности
        """
        m = re.search(r"ВКР (\d+\.\d+\.?\d*)", os.path.basename(self.path))
        if m:
            return list(map(int, m.group(1).split(".")))
        return [0, 1, 0]

    def _next_version_path(self) -> str:
        """Путь с инкрементированной версией (B += 1)."""
        v = self._version.copy()
        if len(v) < 3:
            v.append(1)
        else:
            v[2] += 1
        version_str = ".".join(map(str, v))
        return os.path.join(self._working_dir, f"ВКР {version_str}.docx")

    def set_chapter(self, chapter: int):
        """Переключить на главу: 0.A.0 (сбрасывает правки)."""
        self._version = [0, chapter, 0]

    def mark_ready(self):
        """Пометить ВКР как готовую: 1.0.0."""
        self._version = [1, 0, 0]

    # ============================================================
    # Bookmarks
    # ============================================================

    def _get_max_bookmark_id(self) -> int:
        mx = 0
        for bm in self.doc.element.findall(".//" + qn("w:bookmarkStart")):
            try:
                mx = max(mx, int(bm.get(qn("w:id"), "0")))
            except ValueError:
                pass
        return mx

    def _add_bookmark(self, p_elem, name: str):
        self._bm_counter += 1
        bs = etree.Element(qn("w:bookmarkStart"))
        bs.set(qn("w:id"), str(self._bm_counter))
        bs.set(qn("w:name"), name)
        be = etree.Element(qn("w:bookmarkEnd"))
        be.set(qn("w:id"), str(self._bm_counter))
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(bs)
        else:
            p_elem.insert(0, bs)
        p_elem.append(be)

    # ============================================================
    # XML-хелперы
    # ============================================================

    def _mk_run(self, text: str, bold=False, italic=False, size=None):
        size = size or self._font_size
        r = etree.Element(qn("w:r"))
        rPr = etree.SubElement(r, qn("w:rPr"))
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), self._font)
        if bold:
            etree.SubElement(rPr, qn("w:b")).set(qn("w:val"), "1")
        if italic:
            etree.SubElement(rPr, qn("w:i")).set(qn("w:val"), "1")
        etree.SubElement(rPr, qn("w:color")).set(qn("w:val"), "000000")
        etree.SubElement(rPr, qn("w:sz")).set(qn("w:val"), str(size * 2))
        t = etree.SubElement(r, qn("w:t"))
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text or ""
        return r

    def _set_paragraph_format(self, p, align="both", first_line=None, left=0,
                               space_before=0, space_after=0, line_spacing=None):
        first_line = first_line if first_line is not None else self._first_line_indent
        line_spacing = line_spacing or self._line_spacing

        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = etree.SubElement(p._element, qn("w:pPr"))
            p._element.insert(0, pPr)

        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(pPr, qn("w:jc"))
        jc.set(qn("w:val"), align)

        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = etree.SubElement(pPr, qn("w:ind"))
        ind.set(qn("w:firstLine"), str(first_line))
        ind.set(qn("w:left"), str(left))

        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = etree.SubElement(pPr, qn("w:spacing"))
        sp.set(qn("w:before"), str(space_before))
        sp.set(qn("w:after"), str(space_after))
        sp.set(qn("w:line"), str(line_spacing))
        sp.set(qn("w:lineRule"), "auto")

    # ============================================================
    # Настройка страницы и стилей
    # ============================================================

    def fix_page_setup(self):
        """Поля страницы по ГОСТ."""
        page = self.cfg.get("page", {})
        for section in self.doc.sections:
            section.left_margin = Cm(page.get("margin_left_cm", 3.0))
            section.right_margin = Cm(page.get("margin_right_cm", 1.5))
            section.top_margin = Cm(page.get("margin_top_cm", 2.0))
            section.bottom_margin = Cm(page.get("margin_bottom_cm", 2.0))

    def fix_styles(self):
        """Стили Heading 1/2/3 и Normal по ГОСТ."""
        self._fix_style("Normal", bold=None)
        self._fix_style("Heading 1", bold=True)
        self._fix_style("Heading 2", bold=True)
        self._fix_style("Heading 3", bold=True)

    def _fix_style(self, name, bold=None):
        try:
            style = self.doc.styles[name]
        except KeyError:
            return
        style.font.name = self._font
        style.font.size = Pt(self._font_size)
        if bold is not None:
            style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        self._set_rfont(style.element, self._font)

    def _set_rfont(self, style_elem, font_name):
        rPr = style_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(style_elem, qn("w:rPr"))
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), font_name)
        # Убрать синий цвет из Heading 3
        color = rPr.find(qn("w:color"))
        if color is not None and color.get(qn("w:val"), "") != "000000":
            rPr.remove(color)

    # ============================================================
    # Добавление контента
    # ============================================================

    def add_heading1(self, text: str):
        """Heading 1 — глава. По центру, Bold, с новой страницы."""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Heading 1"]
        run = p.add_run(text)
        run.font.name = self._font
        run.font.size = Pt(self._font_size)
        self._set_paragraph_format(p, align="center", first_line=0, left=0)
        return p

    def add_heading2(self, text: str):
        """Heading 2 — подраздел. JUSTIFY, Bold, отступ 1.25."""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Heading 2"]
        run = p.add_run(text)
        run.font.name = self._font
        run.font.size = Pt(self._font_size)
        self._set_paragraph_format(p, align="both")
        return p

    def add_heading3(self, text: str):
        """Heading 3 — пункт. JUSTIFY, Bold, отступ 1.25."""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles["Heading 3"]
        run = p.add_run(text)
        run.font.name = self._font
        run.font.size = Pt(self._font_size)
        run.bold = True
        self._set_paragraph_format(p, align="both")
        return p

    def add_body_text(self, text: str):
        """Основной текст. JUSTIFY, отступ 1.25, Times New Roman 14pt."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self._font
        run.font.size = Pt(self._font_size)
        self._set_paragraph_format(p, align="both")
        return p

    def add_figure(self, image_path: str, caption: str, width_cm: float = None):
        """Рисунок по центру + подпись. Автомасштабирование по высоте."""
        p_img = self.doc.add_paragraph()
        self._set_paragraph_format(p_img, align="center", first_line=0, left=0)

        if os.path.exists(image_path):
            img = Image.open(image_path)
            w, h = img.size
            aspect = h / w
            if width_cm is None:
                width_cm = self._max_img_w
            target_h = width_cm * aspect
            if target_h > self._max_img_h:
                width_cm = self._max_img_h / aspect
                target_h = self._max_img_h

            run = p_img.add_run()
            run.add_picture(image_path,
                            width=int(width_cm * 360000),
                            height=int(target_h * 360000))
        else:
            p_img.add_run(f"[Файл не найден: {image_path}]")

        p_cap = self.doc.add_paragraph()
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = self._font
        run_cap.font.size = Pt(self._font_size)
        self._set_paragraph_format(p_cap, align="center", first_line=0, left=0)
        return p_img, p_cap

    def add_figure_side_by_side(self, path1: str, path2: str, caption: str,
                                 width_cm: float = None):
        """Два рисунка рядом через таблицу без границ."""
        width_cm = width_cm or self._side_by_side_w
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._remove_table_borders(table)

        for cell, img_path in [(table.cell(0, 0), path1), (table.cell(0, 1), path2)]:
            p = cell.paragraphs[0]
            self._set_paragraph_format(p, align="center", first_line=0, left=0)
            if os.path.exists(img_path):
                img = Image.open(img_path)
                aspect = img.size[1] / img.size[0]
                run = p.add_run()
                run.add_picture(img_path,
                                width=int(width_cm * 360000),
                                height=int(width_cm * aspect * 360000))

        p_cap = self.doc.add_paragraph()
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = self._font
        run_cap.font.size = Pt(self._font_size)
        self._set_paragraph_format(p_cap, align="center", first_line=0, left=0)
        return table, p_cap

    def add_table_caption(self, caption: str):
        """Подпись таблицы — над таблицей, с абзацным отступом."""
        p = self.doc.add_paragraph()
        run = p.add_run(caption)
        run.font.name = self._font
        run.font.size = Pt(self._font_size)
        self._set_paragraph_format(p, align="both")
        return p

    def add_data_table(self, headers: list, rows: list):
        """Таблица с данными. Границы, 12pt, заголовки Bold."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_table_borders(table)

        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(header)
            run.font.name = self._font
            run.font.size = Pt(self._font_size_table)
            run.bold = True
            self._set_paragraph_format(p, align="center", first_line=0, left=0)

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.cell(i + 1, j)
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(val)
                run.font.name = self._font
                run.font.size = Pt(self._font_size_table)
                self._set_paragraph_format(p, align="center", first_line=0, left=0)
        return table

    def _remove_table_borders(self, table):
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.SubElement(table._tbl, qn("w:tblPr"))
        borders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = etree.SubElement(borders, qn(f"w:{name}"))
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")

    def _add_table_borders(self, table):
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.SubElement(table._tbl, qn("w:tblPr"))
        borders = etree.SubElement(tblPr, qn("w:tblBorders"))
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = etree.SubElement(borders, qn(f"w:{name}"))
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")

    # ============================================================
    # Перекрёстные ссылки
    # ============================================================

    def add_crossrefs(self) -> int:
        """Двусторонние перекрёстные ссылки для всех рисунков и таблиц."""
        self._remove_old_crossrefs()
        refs, caps = self._scan_refs_and_caps()

        rPr_cache, text_cache = {}, {}
        for key in set(list(refs.keys()) + list(caps.keys())):
            if key in refs:
                idx = refs[key][0]
                if idx not in rPr_cache:
                    rPr_cache[idx] = self._get_rPr(self.doc.paragraphs[idx])
                    text_cache[idx] = self._safe_text(self.doc.paragraphs[idx])
            if key in caps:
                idx = caps[key]
                if idx not in rPr_cache:
                    rPr_cache[idx] = self._get_rPr(self.doc.paragraphs[idx])

        done = 0
        for key in sorted(set(refs.keys()) & set(caps.keys())):
            ref_idx, search = refs[key]
            cap_idx = caps[key]
            bm_ref, bm_cap = _gen_bookmark_id(), _gen_bookmark_id()

            self._add_bookmark(self.doc.paragraphs[ref_idx]._element, bm_ref)
            self._add_bookmark(self.doc.paragraphs[cap_idx]._element, bm_cap)

            full_text = text_cache.get(ref_idx, self._safe_text(self.doc.paragraphs[ref_idx]))
            self._replace_with_hyperlink(self.doc.paragraphs[ref_idx], full_text, search, bm_cap)
            self._wrap_as_hyperlink(self.doc.paragraphs[cap_idx], bm_ref)
            done += 1
        return done

    def _remove_old_crossrefs(self):
        pattern = re.compile(
            r"(На рисунк|на рисунк|рисунках|в таблиц|Рисунок \d+\s*[—–-]|Таблица \d+\s*[—–-])"
        )
        for hl in list(self.doc.element.findall(".//" + qn("w:hyperlink"))):
            anchor = hl.get(qn("w:anchor"), "")
            if not anchor:
                continue
            hl_text = "".join((t.text or "") for t in hl.findall(".//" + qn("w:t")))
            if pattern.match(hl_text.strip()):
                parent = hl.getparent()
                for child in list(hl):
                    hl.addprevious(child)
                parent.remove(hl)

        body = self.doc.element.body
        for bm in list(body.findall(qn("w:bookmarkStart"))):
            bid = bm.get(qn("w:id"))
            for be in list(body.findall(qn("w:bookmarkEnd"))):
                if be.get(qn("w:id")) == bid:
                    body.remove(be)
            body.remove(bm)

    def _scan_refs_and_caps(self):
        refs, caps = {}, {}
        for i, p in enumerate(self.doc.paragraphs):
            text = self._safe_text(p)
            for m in re.finditer(r"[Нн]а рисунк[еи] (\d+)", text):
                k = ("fig", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, m.group(0))
            for m in re.finditer(r"на рисунке (\d+)", text):
                k = ("fig", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, m.group(0))
            for m in re.finditer(r"рисунках (\d+)", text):
                k = ("fig", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, m.group(0))
            for m in re.finditer(r"в таблиц[еу] (\d+)", text):
                k = ("tab", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, m.group(0))
            for m in re.finditer(r"в таблицах (\d+)", text):
                k = ("tab", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, m.group(0))
            for m in re.finditer(r"в таблиц\w+ \d+\s+и\s+(\d+)", text):
                k = ("tab", int(m.group(1)))
                if k not in refs:
                    refs[k] = (i, f"и {m.group(1)}")

            fc = re.match(r"^Рисунок (\d+)\s*[—–-]", text)
            if fc:
                caps[("fig", int(fc.group(1)))] = i
            tc = re.match(r"^Таблица (\d+)\s*[—–-]", text)
            if tc:
                caps[("tab", int(tc.group(1)))] = i
        return refs, caps

    def _safe_text(self, p):
        return "".join((t.text or "") for t in p._element.findall(".//" + qn("w:t")))

    def _get_rPr(self, p):
        for run in p.runs:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is not None:
                return copy.deepcopy(rPr)
        return None

    def _replace_with_hyperlink(self, p, full_text, search, anchor):
        idx = full_text.find(search)
        if idx < 0:
            return
        for child in list(p._element):
            if child.tag == qn("w:r"):
                p._element.remove(child)
        before = full_text[:idx]
        matched = full_text[idx:idx + len(search)]
        after = full_text[idx + len(search):]

        insert_pt = None
        for child in p._element:
            if child.tag in (qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
                insert_pt = child

        elems = []
        if before:
            elems.append(self._mk_run(before))
        hl = etree.Element(qn("w:hyperlink"))
        hl.set(qn("w:anchor"), anchor)
        hl.append(self._mk_run(matched))
        elems.append(hl)
        if after:
            elems.append(self._mk_run(after))

        if insert_pt is not None:
            for e in reversed(elems):
                insert_pt.addnext(e)

    def _wrap_as_hyperlink(self, p, anchor):
        runs_info = []
        for run in p.runs:
            rPr = run._element.find(qn("w:rPr"))
            runs_info.append({
                "text": run.text or "",
                "rPr": copy.deepcopy(rPr) if rPr is not None else None,
            })
        if not runs_info:
            return
        for child in list(p._element):
            if child.tag == qn("w:r"):
                p._element.remove(child)

        hl = etree.Element(qn("w:hyperlink"))
        hl.set(qn("w:anchor"), anchor)
        for rd in runs_info:
            r = etree.SubElement(hl, qn("w:r"))
            if rd["rPr"] is not None:
                r.append(copy.deepcopy(rd["rPr"]))
            t = etree.SubElement(r, qn("w:t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = rd["text"]

        insert_pt = None
        for child in p._element:
            if child.tag in (qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
                insert_pt = child
        if insert_pt is not None:
            insert_pt.addnext(hl)
        else:
            p._element.append(hl)

    # ============================================================
    # Содержание (TOC)
    # ============================================================

    def update_toc(self) -> int:
        """Обновить содержание по Heading 1/2/3."""
        headings = []
        for i, p in enumerate(self.doc.paragraphs):
            if p.style.name in ("Heading 1", "Heading 2", "Heading 3"):
                bm_name = _gen_bookmark_id()
                self._add_bookmark(p._element, bm_name)
                headings.append({
                    "level": int(p.style.name[-1]),
                    "text": self._safe_text(p),
                    "bookmark": bm_name,
                })

        toc_start = toc_end = None
        for i, p in enumerate(self.doc.paragraphs):
            if self._safe_text(p).strip() == "СОДЕРЖАНИЕ":
                toc_start = i
            if toc_start is not None and i > toc_start and p.style.name == "Heading 1":
                toc_end = i
                break

        if toc_start is None or toc_end is None:
            return 0

        body = self.doc.element.body
        for tp in self.doc.paragraphs[toc_start + 1:toc_end]:
            body.remove(tp._element)

        insert_before = self.doc.paragraphs[toc_start]._element.getnext()
        if insert_before is None:
            return 0

        # Ширина текстового поля в twips
        page = self.cfg.get("page", {})
        text_width_cm = (page.get("width_cm", 21.0)
                         - page.get("margin_left_cm", 3.0)
                         - page.get("margin_right_cm", 1.5))
        tab_pos = str(int(text_width_cm / 2.54 * 1440))

        for h in headings:
            p_elem = etree.Element(qn("w:p"))
            pPr = etree.SubElement(p_elem, qn("w:pPr"))

            tabs = etree.SubElement(pPr, qn("w:tabs"))
            tab = etree.SubElement(tabs, qn("w:tab"))
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), tab_pos)

            ind = etree.SubElement(pPr, qn("w:ind"))
            indent_map = {1: "0", 2: "567", 3: "1134"}
            ind.set(qn("w:left"), indent_map.get(h["level"], "0"))
            ind.set(qn("w:firstLine"), "0")

            sp = etree.SubElement(pPr, qn("w:spacing"))
            sp.set(qn("w:before"), "0")
            sp.set(qn("w:after"), "0")
            sp.set(qn("w:line"), str(self._line_spacing))
            sp.set(qn("w:lineRule"), "auto")

            hl = etree.SubElement(p_elem, qn("w:hyperlink"))
            hl.set(qn("w:anchor"), h["bookmark"])
            hl.append(self._mk_run(h["text"]))

            r_tab = etree.SubElement(hl, qn("w:r"))
            etree.SubElement(r_tab, qn("w:tab"))

            insert_before.addprevious(p_elem)

        return len(headings)

    # ============================================================
    # Ограничение высоты изображений
    # ============================================================

    def fix_image_sizes(self) -> int:
        """Ограничить высоту всех изображений."""
        max_h_emu = int(self._max_img_h * 360000)
        fixed = 0
        ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"

        for drawing in self.doc.element.findall(".//" + qn("w:drawing")):
            inline = drawing.find(f"{{{ns_wp}}}inline")
            if inline is None:
                inline = drawing.find(f"{{{ns_wp}}}anchor")
            if inline is None:
                continue
            extent = inline.find(f"{{{ns_wp}}}extent")
            if extent is None:
                continue

            cx, cy = int(extent.get("cx", 0)), int(extent.get("cy", 0))
            if cy > max_h_emu:
                scale = max_h_emu / cy
                extent.set("cx", str(int(cx * scale)))
                extent.set("cy", str(max_h_emu))
                for ext in inline.findall(f".//{{{ns_a}}}ext"):
                    ext.set("cx", str(int(int(ext.get("cx", 0)) * scale)))
                    ext.set("cy", str(int(int(ext.get("cy", 0)) * scale)))
                fixed += 1
        return fixed

    # ============================================================
    # Сохранение
    # ============================================================

    def save(self, path: str = None) -> str:
        """Сохранить. Без аргумента — инкрементирует версию."""
        if path is None:
            path = self._next_version_path()
        self.doc.save(path)
        return path

    def save_same(self) -> str:
        """Перезаписать текущий файл."""
        self.doc.save(self.path)
        return self.path
